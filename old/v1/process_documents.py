#!/usr/bin/env python3
"""
Document Processor - Extracts UA/UB/UC/UD codes from PDFs and Word files.

Usage:
    python process_documents.py

Reads files from ./input directory, extracts codes matching UA/UB/UC/UD patterns,
and saves results to ./output directory.
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path
from collections import Counter

from PyPDF2 import PdfReader
from docx import Document


# Directories
BASE_DIR = Path(__file__).parent
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"
CONVERTED_DIR = BASE_DIR / "converted"

# Pattern to match UA, UB, UC, UD followed by optional space/dash and one or more digits
CODE_PATTERN = re.compile(r'\b(U[ABCD])\s*[-]?\s*(\d+)\b', re.IGNORECASE)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract all text from a PDF file."""
    text_parts = []
    reader = PdfReader(file_path)

    for page_num, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract all text from a Word document."""
    doc = Document(file_path)
    text_parts = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    return "\n".join(text_parts)


def save_converted_json(file_path: Path, text: str) -> Path:
    """Save extracted text as JSON for intermediate processing."""
    converted_data = {
        "source_file": file_path.name,
        "extracted_at": datetime.now().isoformat(),
        "text_content": text
    }

    output_path = CONVERTED_DIR / f"{file_path.stem}_converted.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(converted_data, f, indent=2, ensure_ascii=False)

    return output_path


def find_codes(text: str) -> dict:
    """Find all UA/UB/UC/UD codes and count occurrences."""
    # Find all matches (case-insensitive, normalized to uppercase)
    # Pattern returns tuples: (prefix, number) e.g., ('UD', '244')
    matches = CODE_PATTERN.findall(text)
    # Combine prefix and number, normalize to uppercase without spaces
    normalized_matches = [f"{prefix.upper()}{num}" for prefix, num in matches]

    # Count occurrences
    code_counts = Counter(normalized_matches)

    # Group by prefix
    grouped = {
        "UA": {},
        "UB": {},
        "UC": {},
        "UD": {}
    }

    for code, count in sorted(code_counts.items()):
        prefix = code[:2]
        if prefix in grouped:
            grouped[prefix][code] = count

    return {
        "all_codes": dict(sorted(code_counts.items())),
        "by_prefix": grouped,
        "total_matches": len(normalized_matches),
        "unique_codes": len(code_counts)
    }


def process_file(file_path: Path) -> dict:
    """Process a single file and return results."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # Save converted JSON
    converted_path = save_converted_json(file_path, text)

    # Find codes
    codes_result = find_codes(text)

    return {
        "source_file": file_path.name,
        "converted_json": converted_path.name,
        "text_length": len(text),
        "codes": codes_result
    }


def main():
    """Main entry point."""
    # Ensure directories exist
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    CONVERTED_DIR.mkdir(exist_ok=True)

    # Find all supported files
    supported_extensions = [".pdf", ".docx", ".doc"]
    files_to_process = [
        f for f in INPUT_DIR.iterdir()
        if f.is_file() and f.suffix.lower() in supported_extensions
    ]

    if not files_to_process:
        print(f"No PDF or Word files found in {INPUT_DIR}")
        print("Supported formats: .pdf, .docx, .doc")
        return

    print(f"Found {len(files_to_process)} file(s) to process")

    # Process each file
    results = {
        "processed_at": datetime.now().isoformat(),
        "files_processed": [],
        "results": [],
        "summary": {
            "total_files": 0,
            "total_codes_found": 0,
            "all_unique_codes": {}
        }
    }

    all_codes_combined = Counter()

    for file_path in files_to_process:
        print(f"Processing: {file_path.name}")
        try:
            file_result = process_file(file_path)
            results["files_processed"].append(file_path.name)
            results["results"].append(file_result)

            # Aggregate codes
            all_codes_combined.update(file_result["codes"]["all_codes"])

            print(f"  - Found {file_result['codes']['total_matches']} code matches")
            print(f"  - Unique codes: {file_result['codes']['unique_codes']}")

        except Exception as e:
            print(f"  - Error processing {file_path.name}: {e}")
            results["results"].append({
                "source_file": file_path.name,
                "error": str(e)
            })

    # Update summary
    results["summary"]["total_files"] = len(results["files_processed"])
    results["summary"]["total_codes_found"] = sum(all_codes_combined.values())
    results["summary"]["all_unique_codes"] = dict(sorted(all_codes_combined.items()))

    # Group summary by prefix
    results["summary"]["by_prefix"] = {
        "UA": {k: v for k, v in all_codes_combined.items() if k.startswith("UA")},
        "UB": {k: v for k, v in all_codes_combined.items() if k.startswith("UB")},
        "UC": {k: v for k, v in all_codes_combined.items() if k.startswith("UC")},
        "UD": {k: v for k, v in all_codes_combined.items() if k.startswith("UD")}
    }

    # Save final results
    output_file = OUTPUT_DIR / f"results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    print(f"\nResults saved to: {output_file}")
    print(f"Converted files saved to: {CONVERTED_DIR}")
    print(f"\nSummary:")
    print(f"  - Files processed: {results['summary']['total_files']}")
    print(f"  - Total code matches: {results['summary']['total_codes_found']}")
    print(f"  - Unique codes: {len(results['summary']['all_unique_codes'])}")


if __name__ == "__main__":
    main()
